'''
input : 
1. image : forecasted image
2. data : forecasted data
3. text : coin_text
4. current date and week day

how to do 
1. batch_coin, day_of_week, current_date = extract_batch_coin()
2. locate download folders
3. load text / image / data file
4. Make word file
5. save word file
'''
import pandas as pd

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.shared import RGBColor

from tqdm import tqdm

import os
from batch_target import extract_batch_coin

'''
FUNCTION
'''

def word_makedirs(day_of_week):
    word_save_path = f"./word data/{str(day_of_week)}/{formatted_date}/"
    # Set the word path
    if not os.path.exists(word_save_path):
        os.makedirs(word_save_path)
    return word_save_path

def auto_word_func():
    # Create a new Word document
    doc = Document()

    header = doc.add_heading(f'{coin_name} Trading Insights: A Comprehensive Financial Report ({week_month_date})', level=1)

    # Access the style of the header
    header_style = header.style
    font = header_style.font
    font.color.rgb = RGBColor(0, 0, 0)
    font.bold = True

    doc.add_paragraph(f'{first_body}')

    # ADD IMAGE
    doc.add_picture(f'{forecasted_image_path}/period/{coin_symbol}_period_{formatted_date}.png', width=Inches(6), height=Inches(4))
    explanation_period = doc.add_paragraph("Explanation: Latest 6 months' price and forecasted 7 days.")

    doc.add_picture(f'{forecasted_image_path}/whole/{coin_symbol}_whole_{formatted_date}.png', width=Inches(6), height=Inches(4))
    explanation_whole = doc.add_paragraph('Explanation: Whole price and forecasted 7 days.')

    # Set the paragraph alignment to center
    explanation_period.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    explanation_whole.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set the font size to 8 for each explanation separately
    font_period = explanation_period.runs[0].font
    font_period.size = Pt(9)

    font_whole = explanation_whole.runs[0].font
    font_whole.size = Pt(9)

    doc.add_paragraph(f'{second_body}')

    doc.save(f'{word_save_path}{coin_name}_trending_report.docx')


# 1. 
batch_coin, day_of_week, current_date = extract_batch_coin()
month, day, week= current_date.strftime("%B"), current_date.strftime("%d"), current_date.strftime("%W")
week_month_date = f"{month} {day}, {current_date.year} (Week {week})"

formatted_date  = current_date.strftime('%Y_%m_%d')

word_save_path = word_makedirs(day_of_week)

# 2. locate download folders
forecasted_data_path  = f'./forecasted data/{str(day_of_week)}/{formatted_date}/' 
forecasted_image_path = f'./forecasted image/{str(day_of_week)}/{formatted_date}/'

coin_text_df = pd.read_csv('./document/coin_text.csv')
model_pargraph = "4. Forecasting Next Week's Price with A.I. Model"

if __name__ == '__main__': 
    # 3. load text / image / data file
    for index, value in tqdm(batch_coin.iterrows(), total=batch_coin.shape[0]):
        coin_name   = value['coin_name']
        coin_symbol = value['symbol'].lower()
        # forcasted information
        forcasted_data  = pd.read_csv(f'{forecasted_data_path}{coin_symbol}_data_{formatted_date}.csv',
                                    index_col=0)
        total_volume    = "{:,.0f}".format(forcasted_data['Total Volume'].iloc[-1])
        market_cap      = "{:,.0f}".format(forcasted_data['Market Cap'].iloc[-1])
        forcasted_data.drop(['Market Cap', 'Total Volume'], axis=1 , inplace=True)

        # coin text template
        body = coin_text_df.loc[coin_text_df['coin_name'] == coin_name, 'template'].values[0]
        start_index = body.find(model_pargraph)
        second_index = body.find("Please note that")

        # To delete useless information
        if start_index != -1 :
            first_body = body[:start_index + len(model_pargraph)]
            first_body = first_body.replace(r'"__today__"', week_month_date)
            first_body = first_body.replace('XXXX', total_volume, 1)
            first_body = first_body.replace('XXXX', market_cap, 2)

            second_body = body[second_index-1:]
        auto_word_func()
